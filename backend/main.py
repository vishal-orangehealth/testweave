import asyncio
import io
import json
import logging
import os
import re
import time
import uuid

import anthropic
import httpx
from dotenv import load_dotenv
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("testweave")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from huggingface_hub import InferenceClient
    HAS_HUGGINGFACE = True
except ImportError:
    HAS_HUGGINGFACE = False

app = FastAPI(title="TestWeave API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── API Configuration ───────────────────────────────────────────────────────

ANTHROPIC_API_KEY    = os.environ.get("ANTHROPIC_API_KEY", "")
HUGGINGFACE_API_TOKEN = os.environ.get("HUGGINGFACE_API_TOKEN", "")

# Initialize Anthropic clients
if ANTHROPIC_API_KEY:
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    async_client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
else:
    client = None
    async_client = None
    print("WARNING: ANTHROPIC_API_KEY not set — Claude will not be available")

# Initialize HuggingFace client if token is available
hf_client = None
if HAS_HUGGINGFACE and HUGGINGFACE_API_TOKEN:
    hf_client = InferenceClient(token=HUGGINGFACE_API_TOKEN)  # FIX: use 'token=' not 'api_key='
elif not HUGGINGFACE_API_TOKEN:
    print("WARNING: HUGGINGFACE_API_TOKEN not set — HuggingFace fallback will not be available")

# ─── Models ──────────────────────────────────────────────────────────────────

class PRDRequest(BaseModel):
    prd_text: str
    project_name: Optional[str] = "Untitled Project"

class FigmaRequest(BaseModel):
    figma_url: str
    figma_token: str
    project_name: Optional[str] = "Untitled Project"

# ─── Prompts ─────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a senior QA engineer with 10+ years of experience writing test cases.
Given a structured description of UI screens or product requirements, generate comprehensive test cases.

Output ONLY valid JSON — no prose, no markdown fences, no explanation.

Schema:
{
  "project": string,
  "source": "PRD" | "Figma",
  "screens": [
    {
      "screen_name": string,
      "test_cases": [
        {
          "id": string,
          "title": string,
          "type": "happy_path" | "edge_case" | "negative" | "accessibility" | "performance",
          "priority": "P0" | "P1" | "P2",
          "preconditions": [string],
          "steps": [string],
          "expected_result": string,
          "component": string
        }
      ]
    }
  ],
  "summary": {
    "total": number,
    "p0": number,
    "p1": number,
    "p2": number,
    "happy_path": number,
    "negative": number,
    "edge_case": number,
    "accessibility": number
  }
}

Rules:
- P0: blocks core user journey (auth, payment, core CRUD)
- P1: important functionality with a workaround
- P2: edge case, cosmetic, nice-to-have
- Always include at least 1 accessibility test per screen
- Always include at least 1 negative/error test per input field found
- IDs format: TC-001, TC-002, ...
- Steps must be user-perspective actions ("Click", "Enter", "Navigate")
- Expected results must be specific and observable
"""

# ─── Figma Extractor ─────────────────────────────────────────────────────────

def parse_figma_file_key(url: str) -> str:
    patterns = [
        r"figma\.com/file/([a-zA-Z0-9]+)",
        r"figma\.com/design/([a-zA-Z0-9]+)",
        r"figma\.com/proto/([a-zA-Z0-9]+)",
    ]
    for p in patterns:
        m = re.search(p, url)
        if m:
            return m.group(1)
    raise ValueError("Could not parse Figma file key from URL")

async def extract_figma_context(file_key: str, token: str) -> dict:
    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.get(
            f"https://api.figma.com/v1/files/{file_key}",
            headers={"X-Figma-Token": token},
            params={"depth": "3"},
        )
        if resp.status_code == 403:
            raise HTTPException(403, "Invalid Figma token or no access to file")
        if resp.status_code != 200:
            raise HTTPException(resp.status_code, f"Figma API error: {resp.text[:200]}")
        data = resp.json()

    screens = []

    def walk(node, depth=0, screen=None):
        t = node.get("type", "")
        name = node.get("name", "")

        if t == "FRAME" and depth == 1:
            screen = {
                "name": name,
                "components": [],
                "texts": [],
                "flows": [],
            }
            screens.append(screen)

        if screen is None:
            for child in node.get("children", []):
                walk(child, depth + 1, screen)
            return

        if t == "INSTANCE":
            props = node.get("componentProperties", {})
            variants = {k: v.get("value") for k, v in props.items()} if props else {}
            screen["components"].append({"name": name, "variants": variants})

        if t == "TEXT":
            content = node.get("characters", "")
            if content.strip():
                screen["texts"].append({"label": name, "content": content[:100]})

        for reaction in node.get("reactions", []):
            action = reaction.get("action", {})
            trigger = reaction.get("trigger", {})
            screen["flows"].append({
                "from": name,
                "trigger": trigger.get("type", ""),
                "action": action.get("type", ""),
            })

        for child in node.get("children", []):
            walk(child, depth + 1, screen)

    doc = data.get("document", {})
    for page in doc.get("children", []):
        for child in page.get("children", []):
            walk(child, depth=1)

    return {
        "file_name": data.get("name", "Figma File"),
        "screens": screens[:10],
    }

# ─── Claude/HuggingFace Generators ───────────────────────────────────────────

async def _call_llm(system: str, user: str, max_tokens: int) -> str:
    """Async LLM call — returns raw text. Raises HTTPException on total failure."""
    anthropic_error = None
    hf_error = None

    if async_client and ANTHROPIC_API_KEY:
        t0 = time.time()
        try:
            msg = await async_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=max_tokens,
                system=system,
                messages=[{"role": "user", "content": user}],
            )
            elapsed = time.time() - t0
            if msg.stop_reason == "max_tokens":
                raise HTTPException(500, f"Response truncated at {max_tokens} tokens — reduce input size or split into smaller chunks")
            log.info(f"Claude responded in {elapsed:.1f}s | tokens_used={msg.usage.output_tokens}")
            return msg.content[0].text.strip()
        except HTTPException:
            raise
        except Exception as e:
            anthropic_error = str(e)
            log.error(f"Anthropic error ({time.time()-t0:.1f}s): {e}")

    if hf_client and HUGGINGFACE_API_TOKEN:
        try:
            response = hf_client.chat_completion(
                model="Qwen/Qwen2.5-72B-Instruct",
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user},
                ],
                max_tokens=max_tokens,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            hf_error = str(e)
            log.error(f"HuggingFace error: {e}")

    if not ANTHROPIC_API_KEY and not HUGGINGFACE_API_TOKEN:
        raise HTTPException(500, "No AI API keys configured.")

    parts = []
    if anthropic_error:
        parts.append(f"Anthropic: {anthropic_error}")
    elif not ANTHROPIC_API_KEY:
        parts.append("Anthropic: ANTHROPIC_API_KEY not set")
    if hf_error:
        parts.append(f"HuggingFace: {hf_error}")
    elif not HUGGINGFACE_API_TOKEN:
        parts.append("HuggingFace: HUGGINGFACE_API_TOKEN not set")
    raise HTTPException(500, " | ".join(parts))


def _parse_json(raw: str):
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*\n?", "", raw)
    raw = re.sub(r"\n?```\s*$", "", raw)
    raw = raw.strip()
    if not raw.startswith(('{', '[')):
        m = re.search(r'[{\[]', raw)
        if m:
            raw = raw[m.start():]
    return json.loads(raw)


SUMMARIZE_PROMPT = """You are a senior product analyst. Given a PRD chunk, extract screens/features mentioned.

Output ONLY a valid JSON array — no prose, no markdown fences.

Schema:
[
  {
    "screen_name": string,
    "input_fields": [string],
    "error_conditions": [string],
    "key_flows": [string]
  }
]

Rules:
- Only include screens explicitly described in this chunk
- Max 5 words per string value
- input_fields: form fields, inputs, dropdowns, toggles
- error_conditions: validation errors, failure states
- key_flows: main user actions on this screen
- If no clear screens found, return []
"""

TESTCASE_PROMPT = """You are a senior QA engineer. Given a screen summary and a specific test type, generate ONLY that type of test cases.

Output ONLY valid JSON array — no prose, no markdown fences.

Schema:
[
  {
    "id": string,
    "title": string,
    "type": string,
    "priority": "P0" | "P1" | "P2",
    "preconditions": [string],
    "steps": [string],
    "expected_result": string,
    "component": string
  }
]

Rules:
- P0: blocks core user journey (auth, payment, core CRUD)
- P1: important functionality with a workaround
- P2: edge case, cosmetic, nice-to-have
- IDs format: TC-001, TC-002, ... (continue from offset given)
- Steps: user-perspective actions ("Click", "Enter", "Navigate")
- Expected results: specific and observable

For each test type, generate exhaustively:
- happy_path: every successful user journey and valid workflow
- negative: every input field × (empty, invalid format, wrong type, boundary exceeded, special chars, SQL injection attempt, XSS attempt)
- edge_case: boundary values, max/min limits, simultaneous actions, race conditions, large data, empty states
- accessibility: keyboard navigation, screen reader labels, focus order, color contrast, ARIA roles, zoom to 200%
- error_states: network failure, timeout, server error, session expiry, permission denied, concurrent edit conflicts
"""

TEST_TYPES = ["happy_path", "negative", "edge_case", "accessibility", "error_states"]


def _chunk_text(text: str, chunk_size: int = 4000, overlap: int = 200) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


def _merge_screens(all_screens: list[list[dict]]) -> list[dict]:
    """Merge screen summaries extracted from multiple chunks, deduplicating by name."""
    merged: dict[str, dict] = {}
    for chunk_screens in all_screens:
        for screen in chunk_screens:
            name = screen.get("screen_name", "").strip()
            if not name:
                continue
            if name not in merged:
                merged[name] = screen
            else:
                existing = merged[name]
                for field in ("input_fields", "error_conditions", "key_flows"):
                    existing_vals = existing.get(field, [])
                    new_vals = screen.get(field, [])
                    combined = list(dict.fromkeys(existing_vals + new_vals))
                    existing[field] = combined
    return list(merged.values())


async def _extract_chunk(i: int, total: int, chunk: str, project_name: str):
    log.info(f"  Chunk {i+1}/{total} — extracting screens ({len(chunk)} chars)...")
    try:
        raw = await _call_llm(
            system=SUMMARIZE_PROMPT,
            user=f"Project: {project_name}\nChunk {i+1}/{total}:\n\n{chunk}",
            max_tokens=800,
        )
        screens = _parse_json(raw)
        if isinstance(screens, list):
            log.info(f"  Chunk {i+1} → found {len(screens)} screen(s)")
            return screens
    except Exception as e:
        log.warning(f"  Chunk {i+1} extraction failed: {e}")
    return []


async def _generate_for_type(screen: dict, test_type: str, tc_offset: int):
    screen_name = screen.get("screen_name", "Unknown Screen")
    log.info(f"    [{screen_name}] generating {test_type}...")
    user_msg = (
        f"Screen: {screen_name}\n"
        f"Test type to generate: {test_type}\n\n"
        f"Screen summary:\n{json.dumps(screen, indent=2)}\n\n"
        f"Start IDs from TC-{tc_offset:03d}. Generate ONLY {test_type} test cases. Be exhaustive."
    )
    try:
        raw = await _call_llm(system=TESTCASE_PROMPT, user=user_msg, max_tokens=2000)
        test_cases = _parse_json(raw)
        if isinstance(test_cases, list):
            log.info(f"    [{screen_name}] {test_type} → {len(test_cases)} cases")
            return test_cases
    except Exception as e:
        log.warning(f"    [{screen_name}] {test_type} failed: {e}")
    return []


async def generate_tests_two_pass(prd_text: str, project_name: str, filename: str = "") -> dict:
    """
    Pass 1: All chunks extracted in PARALLEL → merge screens.
    Pass 2: All screen×type combos generated in PARALLEL.
    """
    total_start = time.time()

    # Pass 1 — all chunks in parallel
    chunks = _chunk_text(prd_text, chunk_size=1500, overlap=100)
    log.info(f"PASS 1 — {len(chunks)} chunks firing in parallel...")

    results = await asyncio.gather(*[
        _extract_chunk(i, len(chunks), chunk, project_name)
        for i, chunk in enumerate(chunks)
    ])
    screens_summary = _merge_screens([r for r in results if r])
    log.info(f"PASS 1 done — {len(screens_summary)} unique screens in {time.time()-total_start:.1f}s")

    if not screens_summary:
        raise HTTPException(500, "Could not extract screens from PRD")

    # Pass 2 — all screen×type combos in parallel
    total_calls = len(screens_summary) * len(TEST_TYPES)
    log.info(f"PASS 2 — {total_calls} calls firing in parallel...")

    tasks = [
        _generate_for_type(screen, test_type, tc_offset=1)
        for screen in screens_summary
        for test_type in TEST_TYPES
    ]
    all_results = await asyncio.gather(*tasks)

    # Reassemble per screen with sequential IDs
    all_screens = []
    tc_counter = 1
    idx = 0
    for screen in screens_summary:
        screen_name = screen.get("screen_name", "Unknown Screen")
        combined = []
        for _ in TEST_TYPES:
            cases = all_results[idx] or []
            for tc in cases:
                tc["id"] = f"TC-{tc_counter:03d}"
                tc_counter += 1
                combined.append(tc)
            idx += 1
        if combined:
            all_screens.append({"screen_name": screen_name, "test_cases": combined})
            log.info(f"  '{screen_name}' — {len(combined)} test cases")

    if not all_screens:
        raise HTTPException(500, "Failed to generate test cases for any screen")

    all_tcs = [tc for s in all_screens for tc in s.get("test_cases", [])]
    log.info(f"DONE — {len(all_tcs)} total test cases in {time.time()-total_start:.1f}s")
    summary = {
        "total": len(all_tcs),
        "p0": sum(1 for t in all_tcs if t.get("priority") == "P0"),
        "p1": sum(1 for t in all_tcs if t.get("priority") == "P1"),
        "p2": sum(1 for t in all_tcs if t.get("priority") == "P2"),
        "happy_path": sum(1 for t in all_tcs if t.get("type") == "happy_path"),
        "negative": sum(1 for t in all_tcs if t.get("type") == "negative"),
        "edge_case": sum(1 for t in all_tcs if t.get("type") == "edge_case"),
        "accessibility": sum(1 for t in all_tcs if t.get("type") == "accessibility"),
        "error_states": sum(1 for t in all_tcs if t.get("type") == "error_states"),
    }

    return {
        "project": project_name,
        "source": "PRD",
        "filename": filename,
        "screens": all_screens,
        "summary": summary,
    }


async def call_claude(user_message: str) -> dict:
    """Single-pass generation (used by Figma)."""
    raw = await _call_llm(system=SYSTEM_PROMPT, user=user_message, max_tokens=16000)
    return _parse_json(raw)

# ─── Document text extractors ────────────────────────────────────────────────

def extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")

def extract_docx(data: bytes) -> str:
    if not HAS_DOCX:
        raise HTTPException(500, "python-docx not installed on server")
    doc = DocxDocument(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading 1" in style:
            parts.append(f"\n# {text}")
        elif "Heading 2" in style:
            parts.append(f"\n## {text}")
        elif "Heading 3" in style:
            parts.append(f"\n### {text}")
        else:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

def extract_pdf(data: bytes) -> str:
    if not HAS_PDF:
        raise HTTPException(500, "PyPDF2 not installed on server")
    reader = PyPDF2.PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)

def extract_text_from_file(filename: str, data: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "docx":
        return extract_docx(data)
    elif ext == "pdf":
        return extract_pdf(data)
    elif ext in ("txt", "md"):
        return extract_txt(data)
    else:
        raise HTTPException(400, f"Unsupported file type: .{ext}. Supported: .docx, .pdf, .txt, .md")

# ─── Routes ──────────────────────────────────────────────────────────────────

# ─── Job Store ───────────────────────────────────────────────────────────────
# In-memory store: job_id → {"status": "processing"|"done"|"error", "result": ..., "error": ...}
jobs: dict = {}


async def _run_prd_job(job_id: str, text: str, project_name: str, filename: str):
    try:
        result = await generate_tests_two_pass(text, project_name, filename=filename)
        jobs[job_id] = {"status": "done", "result": result}
        log.info(f"Job {job_id} completed — {result['summary']['total']} test cases")
    except HTTPException as e:
        msg = str(e.detail) if e.detail else f"HTTP {e.status_code}"
        jobs[job_id] = {"status": "error", "error": msg}
        log.error(f"Job {job_id} failed: {msg}")
    except Exception as e:
        jobs[job_id] = {"status": "error", "error": str(e)}
        log.error(f"Job {job_id} failed: {e}")


@app.get("/health")
def health():
    return {
        "status": "ok",
        "docx": HAS_DOCX,
        "pdf": HAS_PDF,
        "anthropic": bool(ANTHROPIC_API_KEY),
        "huggingface": bool(HUGGINGFACE_API_TOKEN),
        "ai_available": bool(ANTHROPIC_API_KEY or HUGGINGFACE_API_TOKEN),
    }

@app.post("/upload/prd")
async def upload_prd(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    project_name: str = Form("Untitled Project"),
):
    if file.size and file.size > 10 * 1024 * 1024:
        raise HTTPException(400, "File too large (max 10 MB)")

    data = await file.read()
    if not data:
        raise HTTPException(400, "Uploaded file is empty")

    try:
        text = extract_text_from_file(file.filename or "file.txt", data)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(422, f"Could not parse file: {e}")

    if len(text.strip()) < 50:
        raise HTTPException(400, "Extracted text is too short — is the file empty or scanned?")

    job_id = str(uuid.uuid4())
    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(_run_prd_job, job_id, text, project_name, file.filename or "")
    log.info(f"Job {job_id} queued for '{project_name}'")
    return {"job_id": job_id, "status": "processing"}


@app.get("/status/{job_id}")
def get_job_status(job_id: str):
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(404, f"Job {job_id} not found")
    if job["status"] == "error":
        raise HTTPException(500, job["error"])
    return job


@app.post("/generate/prd")
async def generate_from_prd(background_tasks: BackgroundTasks, req: PRDRequest):
    if len(req.prd_text.strip()) < 50:
        raise HTTPException(400, "PRD text is too short")

    job_id = str(uuid.uuid4())
    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(_run_prd_job, job_id, req.prd_text, req.project_name, "")
    log.info(f"Job {job_id} queued for '{req.project_name}'")
    return {"job_id": job_id, "status": "processing"}

@app.post("/generate/figma")
async def generate_from_figma(req: FigmaRequest):
    try:
        file_key = parse_figma_file_key(req.figma_url)
    except ValueError as e:
        raise HTTPException(400, str(e))

    context = await extract_figma_context(file_key, req.figma_token)

    screens_summary = []
    for s in context["screens"]:
        screens_summary.append({
            "screen": s["name"],
            "components": s["components"][:15],
            "text_content": s["texts"][:20],
            "user_flows": s["flows"][:10],
        })

    user_msg = f"""Project: {req.project_name}
Source: Figma design file "{context['file_name']}"

Screens found in design:
{json.dumps(screens_summary, indent=2)}

Instructions:
1. For each screen, analyse the components (buttons, inputs, dropdowns) and their variants
2. Use text content (labels, placeholders, error messages) to understand validation rules
3. Use user flows (prototype links) to understand navigation paths
4. Generate test cases: happy path per screen, all component variants, all visible error states,
   navigation flows, and accessibility checks
5. Component variant names (e.g. "Disabled", "Error", "Loading") = separate test cases
"""
    result = await call_claude(user_msg)
    result["source"] = "Figma"
    result["project"] = req.project_name
    return result